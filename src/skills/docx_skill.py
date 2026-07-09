from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


class DocxSkill:
    name = "docx"

    def _set_default_font(self, document: Document):
        styles = document.styles

        normal_style = styles["Normal"]
        normal_style.font.name = "標楷體"
        normal_style.font.size = Pt(12)
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    def _add_text(self, document: Document, text: str):
        for line in text.splitlines():
            line = line.strip()

            if not line:
                document.add_paragraph("")
                continue

            paragraph = document.add_paragraph(line)

            for run in paragraph.runs:
                run.font.name = "標楷體"
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    def _write_morning_docx(self, text: str, output_path: str):
        document = Document()
        self._set_default_font(document)

        title = document.add_heading("每日金融晨報", level=1)
        for run in title.runs:
            run.font.name = "標楷體"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

        self._add_text(document, text)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    def _write_signoff_docx(self, text: str, output_path: str):
        document = Document()
        self._set_default_font(document)

        self._add_text(document, text)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    def run(self, inputs: dict, step_config: dict) -> dict:
        text = inputs["text"]
        output_path = inputs["output_path"]
        mode = step_config.get("mode", "morning")

        if mode == "morning":
            self._write_morning_docx(text, output_path)

        elif mode == "signoff":
            self._write_signoff_docx(text, output_path)

        else:
            raise ValueError(f"不支援的 docx mode：{mode}")

        return {
            "path": output_path
        }
